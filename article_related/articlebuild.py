import os

from article_related.article import article
from config.config import readpathdir
from docx import Document
from input.readsentences import zykeywordextract, menuextract
from input.rwsinput import rws

def article_build():

    files=os.listdir(readpathdir)
    articlelist=[]
    for file in files:
        print(file)
        file_path = os.path.join(readpathdir, file)
        doc = Document(file_path)
        print(file_path)
        print(len(doc.paragraphs))
        name = file.split('.')[0]
        x=article()
        x.setname(name)
        x.setparagraphs(doc.paragraphs)
        articlelist.append(x)
    #rws build
    for a in articlelist:
        a.setrws(rws(a.getparagraphs(),a.getname()))#������
        res=zykeywordextract(a.getparagraphs(),a)#ժҪ�͹ؼ���
        a.setzy(res[0])
        a.setkeyword(res[1])
        a.setmenu(menuextract(a.getparagraphs(),a))#Ŀ¼
    return articlelist
articlelist=article_build()
m=1